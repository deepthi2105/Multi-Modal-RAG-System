import os
from dotenv import load_dotenv
import streamlit as st
import PyPDF2
import docx
import whisper
import numpy as np
import faiss
from PIL import Image
import pytesseract
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.docstore.document import Document
from openai import AzureOpenAI

load_dotenv()
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2023-12-01-preview",
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
)

# ---------------------- Document Extraction ----------------------
def extract_documents_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    docs = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            docs.append(Document(page_content=text, metadata={"source": file.name, "page": page_num + 1}))
    return docs

def extract_documents_from_docx(file):
    doc = docx.Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [Document(page_content=full_text, metadata={"source": file.name})]

def extract_documents_from_txt(file):
    text = file.read().decode('utf-8')
    return [Document(page_content=text, metadata={"source": file.name})]

def extract_documents_from_image(file):
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    return [Document(page_content=text, metadata={"source": file.name})]

def extract_documents_from_video(file):
    whisper_model = whisper.load_model("base")
    with open("temp_video.mp4", "wb") as f:
        f.write(file.getbuffer())
    result = whisper_model.transcribe("temp_video.mp4", word_timestamps=True)
    docs = []
    for segment in result.get('segments', []):
        docs.append(Document(
            page_content=segment['text'],
            metadata={"source": file.name, "start_time": segment['start'], "end_time": segment['end']}
        ))
    return docs

def extract_documents(file):
    ext = file.name.split(".")[-1].lower()
    if ext == "pdf": return extract_documents_from_pdf(file)
    if ext == "docx": return extract_documents_from_docx(file)
    if ext == "txt": return extract_documents_from_txt(file)
    if ext in ["jpg", "jpeg", "png"]: return extract_documents_from_image(file)
    if ext == "mp4": return extract_documents_from_video(file)
    return []

# ---------------------- Embedding and Chunking ----------------------
def split_documents(documents, chunk_size=500, chunk_overlap=50):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(documents)

def embed_documents_openai(documents):
    embed_model = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
    embeddings = embed_model.embed_documents([doc.page_content for doc in documents])
    return np.array(embeddings)

def build_faiss_index(embeddings):
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)
    return index

# ---------------------- Retrieval and QA ----------------------
def retrieve_documents(query, documents, top_k=5):
    embed_model = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
    query_embedding = np.array(embed_model.embed_query(query)).reshape(1, -1)
    embeddings = embed_model.embed_documents([doc.page_content for doc in documents])
    embeddings = np.array(embeddings)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)
    _, indices = index.search(query_embedding, min(top_k, len(documents)))
    return [documents[i] for i in indices[0]]

def ask_llm(context, question):
    prompt = f"""
You are a helpful assistant. Use only the given context to answer the user's question. Do not use prior knowledge.

If the answer is not available in the context, respond with:
"Answer not found in the provided content."

Context:
{context}

Question: {question}
"""
    response = client.chat.completions.create(
        model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        messages=[{"role": "system", "content": prompt}],
        temperature=0.2,
        max_tokens=700
    )
    return response.choices[0].message.content

def summarize_all(documents):
    if not documents:
        return "No content to summarize."
    full_text = "\n\n".join([doc.page_content for doc in documents])
    prompt = f"""
You are an assistant tasked with summarizing the following content into a clear and concise overview.

Content:
{full_text}

Please provide a detailed summary.
"""
    response = client.chat.completions.create(
        model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        messages=[{"role": "system", "content": prompt}],
        temperature=0.3,
        max_tokens=1000
    )
    return response.choices[0].message.content

# ---------------------- Streamlit UI ----------------------
st.set_page_config(page_title="Multimodal RAG App", layout="wide")
st.title("📄️🎥️🖼️ Multimodal RAG App")
st.caption("Upload PDFs, DOCX, TXT, MP4, JPG, PNG ➔ Retrieve + Trace Sources (Page/Time)")

if 'documents' not in st.session_state:
    st.session_state['documents'] = []

uploaded_files = st.file_uploader(
    "Upload your documents/videos/images", 
    type=["pdf", "docx", "txt", "mp4", "jpg", "jpeg", "png"],
    accept_multiple_files=True
)

if uploaded_files and st.button("Process Files"):
    all_docs = []
    for file in uploaded_files:
        docs = extract_documents(file)
        all_docs.extend(docs)
    st.session_state['documents'] = split_documents(all_docs)
    st.success(f"Processed {len(st.session_state['documents'])} chunks from {len(uploaded_files)} files!")

if st.session_state['documents']:
    if st.button("Summarize All Files"):
        summary = summarize_all(st.session_state['documents'])
        st.subheader("📁 Summary:")
        st.success(summary)

    query = st.text_input("Ask your question:")
    if query:
        results = retrieve_documents(query, st.session_state['documents'])
        if results:
            context = "\n\n".join([doc.page_content for doc in results])
            answer = ask_llm(context, query)
            st.subheader("✅ Answer:")
            st.success(answer)

            st.subheader("📚 Sources Used:")
            seen = set()
            for doc in results:
                src = doc.metadata.get("source", "Unknown")
                page = doc.metadata.get("page")
                start = doc.metadata.get("start_time")
                end = doc.metadata.get("end_time")
                key = (src, page, start, end)
                if key not in seen:
                    seen.add(key)
                    if page:
                        st.markdown(f"- {src} (Page {page})")
                    elif start and end:
                        st.markdown(f"- {src} (from {start:.2f}s to {end:.2f}s)")
                    else:
                        st.markdown(f"- {src}")
        else:
            st.warning("⚠️ No relevant documents found.")
